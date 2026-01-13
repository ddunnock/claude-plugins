"""Document ingestion for PDF, DOCX, HTML, and Markdown.

This module provides:
- Text extraction from various document formats
- Intelligent chunking with overlap
- Optional image extraction (PDF only)
- Metadata extraction (page count, word count, etc.)
"""

import hashlib
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# Optional dependencies
try:
    import fitz as pymupdf  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


@dataclass
class DocumentChunk:
    """A chunk of document content."""

    content: str
    chunk_index: int = 0
    page_number: Optional[int] = None
    section: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "chunk_index": self.chunk_index,
            "page_number": self.page_number,
            "section": self.section,
            "metadata": self.metadata
        }


@dataclass
class IngestedDocument:
    """Metadata about an ingested document."""

    id: str
    filename: str
    file_type: str
    file_hash: str
    file_path: str
    chunk_count: int
    word_count: int
    page_count: Optional[int] = None
    title: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "filename": self.filename,
            "file_type": self.file_type,
            "file_hash": self.file_hash,
            "file_path": self.file_path,
            "chunk_count": self.chunk_count,
            "word_count": self.word_count,
            "page_count": self.page_count,
            "title": self.title,
            "metadata": self.metadata
        }


class DocumentIngestor:
    """Ingest documents and create session events from chunks."""

    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".html": "html",
        ".htm": "html",
        ".md": "markdown",
        ".txt": "text"
    }

    def __init__(self):
        self.available_parsers = {
            "pdf": PYMUPDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "html": BS4_AVAILABLE,
            "markdown": True,  # Built-in
            "text": True  # Built-in
        }

    def get_available_formats(self) -> List[str]:
        """Get list of available document formats."""
        return [fmt for fmt, available in self.available_parsers.items() if available]

    def get_missing_dependencies(self) -> Dict[str, str]:
        """Get missing dependencies for unavailable formats."""
        deps = {
            "pdf": "PyMuPDF (pip install pymupdf)",
            "docx": "python-docx (pip install python-docx)",
            "html": "BeautifulSoup4 (pip install beautifulsoup4)"
        }
        return {fmt: dep for fmt, dep in deps.items() if not self.available_parsers.get(fmt)}

    def ingest(
        self,
        file_path: str,
        chunk_size: int = 1000,
        overlap: int = 200,
        extract_images: bool = False
    ) -> Tuple[List[DocumentChunk], IngestedDocument]:
        """
        Ingest a document and return chunks with metadata.

        Args:
            file_path: Path to the document file
            chunk_size: Maximum characters per chunk
            overlap: Characters to overlap between chunks
            extract_images: Whether to extract images (PDF only)

        Returns:
            Tuple of (list of chunks, document metadata)
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {list(self.SUPPORTED_TYPES.keys())}")

        file_type = self.SUPPORTED_TYPES[ext]
        if not self.available_parsers.get(file_type):
            missing = self.get_missing_dependencies().get(file_type, "unknown dependency")
            raise RuntimeError(f"Parser not available for {file_type}. Install: {missing}")

        # Calculate file hash for deduplication
        with open(path, "rb") as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()[:16]

        # Extract text based on type
        if file_type == "pdf":
            text, metadata = self._extract_pdf(path, extract_images)
        elif file_type == "docx":
            text, metadata = self._extract_docx(path)
        elif file_type == "html":
            text, metadata = self._extract_html(path)
        elif file_type == "markdown":
            text, metadata = self._extract_markdown(path)
        else:  # text
            text, metadata = self._extract_text(path)

        # Count words
        word_count = len(text.split())

        # Chunk the text
        chunks = self._chunk_text(
            text,
            chunk_size,
            overlap,
            metadata.get("page_breaks", []),
            metadata.get("sections", [])
        )

        # Build document metadata
        doc = IngestedDocument(
            id=f"doc-{file_hash}",
            filename=path.name,
            file_type=file_type,
            file_hash=file_hash,
            file_path=str(path.absolute()),
            chunk_count=len(chunks),
            word_count=word_count,
            page_count=metadata.get("page_count"),
            title=metadata.get("title"),
            metadata=metadata
        )

        return chunks, doc

    def _extract_pdf(self, path: Path, extract_images: bool) -> Tuple[str, Dict]:
        """Extract text from PDF using PyMuPDF."""
        doc = pymupdf.open(path)
        text_parts = []
        page_breaks = []
        images = []

        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            text_parts.append(page_text)
            page_breaks.append(len("\n\n".join(text_parts)))

            # Extract images if requested
            if extract_images:
                for img_index, img in enumerate(page.get_images()):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        images.append({
                            "page": page_num + 1,
                            "index": img_index,
                            "ext": base_image["ext"],
                            "size": len(base_image["image"])
                        })
                    except Exception:
                        pass

        doc.close()

        return "\n\n".join(text_parts), {
            "page_count": len(text_parts),
            "page_breaks": page_breaks,
            "images": images if extract_images else []
        }

    def _extract_docx(self, path: Path) -> Tuple[str, Dict]:
        """Extract text from DOCX."""
        doc = DocxDocument(path)

        paragraphs = []
        sections = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings
            if para.style and para.style.name.startswith("Heading"):
                current_section = text
                sections.append({
                    "title": text,
                    "level": int(para.style.name.replace("Heading ", "")) if "Heading " in para.style.name else 1,
                    "position": len("\n\n".join(paragraphs))
                })

            paragraphs.append(text)

        return "\n\n".join(paragraphs), {
            "paragraph_count": len(paragraphs),
            "sections": sections
        }

    def _extract_html(self, path: Path) -> Tuple[str, Dict]:
        """Extract text from HTML."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        soup = BeautifulSoup(content, "html.parser")

        # Extract title
        title = None
        if soup.title:
            title = soup.title.string

        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        # Extract text with some structure preservation
        text_parts = []

        # Process main content areas first
        main_content = soup.find("main") or soup.find("article") or soup.find("body")
        if main_content:
            for element in main_content.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td"]):
                text = element.get_text(separator=" ", strip=True)
                if text:
                    text_parts.append(text)

        # Fallback to simple text extraction
        if not text_parts:
            text = soup.get_text(separator="\n")
            text_parts = [line.strip() for line in text.split("\n") if line.strip()]

        return "\n\n".join(text_parts), {
            "title": title
        }

    def _extract_markdown(self, path: Path) -> Tuple[str, Dict]:
        """Extract text from Markdown file."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        # Extract sections from headers
        sections = []
        lines = text.split("\n")

        for i, line in enumerate(lines):
            if line.startswith("#"):
                # Count heading level
                level = len(line) - len(line.lstrip("#"))
                title = line.lstrip("#").strip()
                if title:
                    sections.append({
                        "title": title,
                        "level": level,
                        "line": i
                    })

        # Extract title from first H1
        title = None
        for section in sections:
            if section["level"] == 1:
                title = section["title"]
                break

        return text, {
            "title": title,
            "sections": sections
        }

    def _extract_text(self, path: Path) -> Tuple[str, Dict]:
        """Extract text from plain text file."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        return text, {}

    def _chunk_text(
        self,
        text: str,
        chunk_size: int,
        overlap: int,
        page_breaks: List[int],
        sections: List[Dict]
    ) -> List[DocumentChunk]:
        """Split text into overlapping chunks with metadata."""
        chunks = []
        start = 0
        chunk_idx = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at natural boundaries
            if end < len(text):
                # Priority: paragraph break, sentence end, word boundary
                for sep in ["\n\n", ".\n", ". ", "\n", " "]:
                    break_point = text.rfind(sep, start + chunk_size // 2, end + 100)
                    if break_point > start:
                        end = break_point + len(sep)
                        break

            chunk_text = text[start:end].strip()

            if chunk_text:
                # Determine page number
                page_num = None
                for i, pb in enumerate(page_breaks):
                    if start < pb:
                        page_num = i + 1
                        break

                # Determine section
                current_section = None
                for section in sections:
                    section_pos = section.get("position") or section.get("line", 0) * 80
                    if section_pos <= start:
                        current_section = section["title"]
                    else:
                        break

                chunks.append(DocumentChunk(
                    content=chunk_text,
                    chunk_index=chunk_idx,
                    page_number=page_num,
                    section=current_section,
                    metadata={
                        "start_offset": start,
                        "end_offset": end,
                        "char_count": len(chunk_text)
                    }
                ))
                chunk_idx += 1

            # Move to next chunk with overlap
            start = max(start + 1, end - overlap)

        return chunks

    def create_events_from_chunks(
        self,
        chunks: List[DocumentChunk],
        doc: IngestedDocument,
        category: str = "document",
        event_type: str = "chunk"
    ) -> List[Dict[str, Any]]:
        """
        Create event payloads from document chunks.

        Args:
            chunks: List of document chunks
            doc: Document metadata
            category: Event category for the chunks
            event_type: Event type for the chunks

        Returns:
            List of event data dictionaries ready for session_record()
        """
        events = []

        for chunk in chunks:
            event_data = {
                "document_id": doc.id,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "chunk_index": chunk.chunk_index,
                "total_chunks": doc.chunk_count,
                "content": chunk.content,
                "page_number": chunk.page_number,
                "section": chunk.section,
                "word_count": len(chunk.content.split())
            }
            events.append(event_data)

        return events
